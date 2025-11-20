# gerar_teste_v43.py
# Gera Teste Diário — M23 (Versão 43) em .docx (Enunciado + Soluções)
# Requisitos: pip install python-docx

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# --- Configuração: caminho de saída (Desktop do utilizador) ---
out_filename = "Teste_Diario_M23_V43_Completo.docx"
out_path = os.path.join(os.path.expanduser("~/Desktop"), out_filename)

# --- Cria documento ---
doc = Document()
doc.styles['Normal'].font.name = 'Helvetica'
doc.styles['Normal'].font.size = Pt(11)

# Cabeçalho / título
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Teste Diário — M23 Engenharia Informática — Versão 43 (Completo)")
r.bold = True
r.font.size = Pt(14)

doc.add_paragraph("")  # espaço

# ---------- ENUNCIADO (todas as 23 questões) ----------
doc.add_paragraph("ENUNCIADO\n").runs[0].bold = True

# Q1
doc.add_paragraph("1) Identidade Combinatória (diferença — sem simetria feita)")
doc.add_paragraph("Resolve em n:")
doc.add_paragraph("C(n+3, 3) – C(n+2, 3) = 12(n+1)")

# Q2
doc.add_paragraph("\n2) Probabilidades — baralho de 52 cartas")
doc.add_paragraph("Tiram-se 2 cartas sem reposição.")
doc.add_paragraph("a) P(2 damas)")
doc.add_paragraph("b) P(2 cartas de naipes diferentes)")
doc.add_paragraph("c) P(1 figura OU carta preta)")

# Q3 Estatística (tabela com dados salteados; N=180)
doc.add_paragraph("\n3) Estatística — Tabela (N = 180)")
doc.add_paragraph("Preenche todos os campos e calcula média, mediana e moda.")
table = doc.add_table(rows=6, cols=5)
hdr = table.rows[0].cells
hdr[0].text = "Classe"
hdr[1].text = "nᵢ"
hdr[2].text = "fᵢ"
hdr[3].text = "Nᵢ"
hdr[4].text = "Fᵢ"

# Linhas: alguns valores fornecidos, outros em branco para preencher
data = [
    ["1", "28", "", "", "0,16"],
    ["2", "", "0,22", "", ""],
    ["3", "46", "", "", ""],
    ["4", "", "0,18", "", ""],
    ["Soma", "180", "", "", "1"]
]
for i, row in enumerate(data):
    cells = table.rows[i+1].cells
    for j, val in enumerate(row):
        cells[j].text = val

doc.add_paragraph("Observações: última classe ajusta para N=180.")

# Q4
doc.add_paragraph("\n4) Probabilidades (dados e cartas)")
doc.add_paragraph("a) 2 dados: P(soma = 9)")
doc.add_paragraph("b) 2 dados: P(pelo menos um 4)")
doc.add_paragraph("c) Baralho: P(preta OU figura)")

# Q5
doc.add_paragraph("\n5) Sucessão racional")
doc.add_paragraph("a_n = (2n + 7) / (n + 5)")
doc.add_paragraph("a) Calcula a_{n+1}")
doc.add_paragraph("b) Estuda o sinal de (a_{n+1} - a_n)")
doc.add_paragraph("c) Limite")
doc.add_paragraph("d) Classificação")

# Q6
doc.add_paragraph("\n6) Função logarítmica")
doc.add_paragraph("f(x) = ln( 1 / (e^x – 2) )")
doc.add_paragraph("a) Domínio")
doc.add_paragraph("b) Zeros")

# Q7
doc.add_paragraph("\n7) Radical e |cos x|")
doc.add_paragraph("g(x) = √((x-2)(x-4)) / (1 – |cos x|)")
doc.add_paragraph("Domínio")
doc.add_paragraph("Zeros")

# Q8
doc.add_paragraph("\n8) Função por ramos")
doc.add_paragraph("f(x) =")
doc.add_paragraph("    ln(1 - x)   se x > 0")
doc.add_paragraph("    e^x - 1     se x ≤ 0")
doc.add_paragraph("Estudar continuidade em x = 0.")

# Q9
doc.add_paragraph("\n9) Função quadrática")
doc.add_paragraph("f(x) = x^2 − 6x + 5")
doc.add_paragraph("a) Zeros")
doc.add_paragraph("b) Vértice + eixo de simetria")
doc.add_paragraph("c) Concavidade")

# Q10
doc.add_paragraph("\n10) Limites — forma 1")
doc.add_paragraph("L1 = lim (n→∞) [ √(n^2 + 8n + 1) − n ]")
doc.add_paragraph("L2 = lim (n→∞) [ √(n^2 + 12n + 16) − √(n^2 + 4n + 4) ]")
doc.add_paragraph("L3 = lim (n→∞) [ √(3n + 4) − √(3n + 1) ]")

# Q11
doc.add_paragraph("\n11) Quadrado Perfeito")
doc.add_paragraph("lim (x→2) (x^2 − 4x + 4) / (x^2 − 2x)")

# Q12
doc.add_paragraph("\n12) Diferença de Cubos")
doc.add_paragraph("lim (x→1) (x^3 − 1) / (x^4 − 1)")

# Q13
doc.add_paragraph("\n13) Grau 2 / Grau 4")
doc.add_paragraph("lim (x→∞) (4x^2 − x + 1) / (x^4 + x^3 − 2)")

# Q14
doc.add_paragraph("\n14) Derivadas básicas")
doc.add_paragraph("a) 5x^4")
doc.add_paragraph("b) x^2 − 3x")
doc.add_paragraph("c) 7/x^3")
doc.add_paragraph("d) √(3x + 2)")

# Q15
doc.add_paragraph("\n15) Derivada pela definição")
doc.add_paragraph("f(x) = x^2")
doc.add_paragraph("Calcular f'(x) pela definição.")

# Q16
doc.add_paragraph("\n16) Derivadas (cadeia, produto, quociente)")
doc.add_paragraph("a) e^(x cos x)")
doc.add_paragraph("b) (sin(2x)) / x^2")

# Q17
doc.add_paragraph("\n17) Produto")
doc.add_paragraph("h(x) = x^3 cos x")

# Q18
doc.add_paragraph("\n18) Quociente")
doc.add_paragraph("g(x) = x^2 / (x + 1)")

# Q19
doc.add_paragraph("\n19) Regra da Cadeia")
doc.add_paragraph("k(x) = √(5x + 1)")

# Q20
doc.add_paragraph("\n20) Derivada + tangente")
doc.add_paragraph("f(x) = x^2 e^x")
doc.add_paragraph("Tangente no ponto x = 1")

# Q21
doc.add_paragraph("\n21) Limite racional")
doc.add_paragraph("lim (x→−1) (x^2 − 1) / (x^3 + 1)")

# Q22
doc.add_paragraph("\n22) Limite trigonométrico")
doc.add_paragraph("lim (x→0) tan x / x")

# Q23
doc.add_paragraph("\n23) Limite trigonométrico (1 − cos)")
doc.add_paragraph("lim (x→0) (1 − cos(5x)) / x")

# ---------- Página de soluções ----------
doc.add_page_break()
doc.add_paragraph("SOLUÇÕES — Versão 43\n").runs[0].bold = True

sols = [
    "Q1 — C(n+3,3) − C(n+2,3) = C(n+2,2) ⇒ (n+2)(n+1)/2 = 12(n+1) ⇒ n = 22.",
    "Q2a — P(2 damas) = C(4,2) / C(52,2) = 6/1326 = 1/221.",
    "Q2b — P(2 naipes diferentes) = (52*39)/(52*51) = (C(4,2)*13*13)/C(52,2) = 13/17. (ver demonstração no exercício)",
    "Q2c — P(1 figura OU carta preta) = (12 + 26 − 6)/52 = 32/52 = 8/13.",
    "Q3 — Tabela (N=180). Exemplo de ni compatíveis: 28, 40, 46, 66. Calcula f_i = n_i/180; N_i acumulada; F_i = N_i/180. Média: (1*28 + 2*40 + 3*46 + 4*66)/180 ≈ 2.95. Moda: classe 3. Mediana: classe 3.",
    "Q4a — P(soma=9) em dois dados = 4/36 = 1/9.",
    "Q4b — P(pelo menos um 4) = 1 − (5/6)^2 = 11/36.",
    "Q4c — P(preta OU figura) = (26 + 12 − 6)/52 = 32/52 = 8/13.",
    "Q5 — a_{n+1} = (2(n+1)+7)/(n+1+5) = (2n+9)/(n+6). A sequência tem limite 2; estudar sinal da diferença dá comportamento crescente para n grandes.",
    "Q6 — f(x)=ln(1/(e^x−2)). Domínio: e^x − 2 > 0 ⇒ x > ln 2. Zeros: 1/(e^x−2)=1 ⇒ e^x = 3 ⇒ x = ln 3.",
    "Q7 — g(x)=√((x−2)(x−4))/(1 − |cos x|). Domínio: x ≤ 2 ou x ≥ 4, e |cos x| ≠ 1. Zeros: x = 2, x = 4.",
    "Q8 — Limites laterais em 0: ln(1−x) → 0 quando x→0^+, e^x−1 → 0 quando x→0^−. Portanto contínua em 0.",
    "Q9 — f(x)=x^2 − 6x + 5 ⇒ zeros em x=1 e x=5. Vértice em x_v = 3; y_v = −4. Concavidade para cima (a>0).",
    "Q10 — L1 = 4; L2 = 4; L3 = 0 (usar racionalização/conjugados).",
    "Q11 — lim (x→2) (x^2−4x+4)/(x^2−2x) = lim ((x−2)^2)/(x(x−2)) = lim (x−2)/x = 0.",
    "Q12 — lim (x→1) (x^3−1)/(x^4−1) = 3/4 (fatorizando e simplificando).",
    "Q13 — lim (x→∞) (4x^2−x+1)/(x^4+x^3−2) = 0 (grau inferior/ superior).",
    "Q14 — Derivadas: a) 20x^3; b) 2x−3; c) −21/x^4; d) 3/(2√(3x+2)).",
    "Q15 — Pela definição f'(x) = lim_{h→0} ((x+h)^2−x^2)/h = 2x.",
    "Q16a — (e^{x cos x})' = e^{x cos x} (cos x − x sin x).",
    "Q16b — (sin(2x)/x^2)' = 2(x cos(2x) − sin(2x))/x^3 (simplificado).",
    "Q17 — (x^3 cos x)' = 3x^2 cos x − x^3 sin x.",
    "Q18 — (x^2/(x+1))' = (x^2 + 2x)/(x+1)^2.",
    "Q19 — d/dx √(5x+1) = 5/(2√(5x+1)).",
    "Q20 — f(x)=x^2 e^x ⇒ f'(x)=e^x(x^2+2x). Em x=1: f(1)=e, f'(1)=3e. Tangente: y = 3e x − 2e.",
    "Q21 — lim (x→−1) (x^2−1)/(x^3+1) = −2/3.",
    "Q22 — lim (x→0) tan x / x = 1.",
    "Q23 — lim (x→0) (1 − cos(5x))/x = 0."
]

for s in sols:
    p = doc.add_paragraph(s)
    p.paragraph_format.space_after = Pt(4)

# --- Grava o ficheiro no Desktop ---
doc.save(out_path)
print("✅ Documento gerado em:", out_path)