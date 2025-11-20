# gera_m23_v43.py
# Gera "Teste_Diario_M23_V43_Completo.docx" no Desktop do utilizador ruiferreira
# Requisitos: python-docx (pip install python-docx)

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_normal_font(doc, name='Helvetica', size=11):
    style = doc.styles['Normal']
    style.font.name = name
    # compatibilidade com algumas versões do python-docx
    try:
        style._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    except Exception:
        pass
    style.font.size = Pt(size)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)

def add_header_paragraph(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)

def add_question(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)

def add_solution_line(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(2)

def main():
    # ajusta aqui se quiseres outro caminho/nome
    desktop_path = "/Users/ruiferreira/Desktop"
    filename = "Teste_Diario_M23_V43_Completo.docx"
    fullpath = f"{desktop_path}/{filename}"

    doc = Document()
    set_normal_font(doc, name='Helvetica', size=11)

    # Title
    add_title(doc, "Teste Diário — M23 Engenharia Informática — Versão 43 (Completo)")
    doc.add_paragraph("")  # espaço

    # ENUNCIADO (toda a Parte A)
    add_header_paragraph(doc, "ENUNCIADO")
    # 1
    add_question(doc, "1) Identidade Combinatória\nResolve em n:\nC(n+3, 3) – C(n+2, 3) = 12(n+1)")
    # 2
    add_question(doc, "2) Probabilidades — baralho de 52 cartas\nTiram-se 2 cartas sem reposição.\na) P(2 damas)\nb) P(2 cartas de naipes diferentes)\nc) P(1 figura OU carta preta)")
    # 3 - tabela salteada (N=180)
    add_question(doc, "3) Estatística — Tabela (N = 180)\nPreenche todos os campos e calcula média, mediana e moda.\n")
    table = doc.add_table(rows=6, cols=5)
    hdr = table.rows[0].cells
    headers = ["Classe", "nᵢ", "fᵢ", "Nᵢ", "Fᵢ"]
    for i, h in enumerate(headers):
        hdr[i].text = h
    # dados salteados (exemplo)
    data = [
        ["1", "28", "", "", ""],
        ["2", "", "0,22", "", ""],
        ["3", "46", "", "", ""],
        ["4", "", "0,18", "", ""],
        ["Soma", "180", "", "", "1"]
    ]
    for i, row in enumerate(data):
        cells = table.rows[i+1].cells
        for j, val in enumerate(row):
            cells[j].text = val
    # centragem
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Observação: última classe ajusta para N=180.")
    # 4
    add_question(doc, "4) Probabilidades (dados e cartas)\na) 2 dados: P(soma = 9)\nb) 2 dados: P(pelo menos um 4)\nc) Baralho: P(preta OU figura)")
    # 5
    add_question(doc, "5) Sucessão racional\n a_n = (2n + 7)/(n + 5)\na) Calcula a_{n+1}\nb) Estuda sinal de (a_{n+1} – a_n)\nc) Limite\nd) Classificação")
    # 6
    add_question(doc, "6) Função logarítmica\nf(x) = ln( 1 / (e^x – 2) )\na) Domínio\nb) Zeros")
    # 7
    add_question(doc, "7) Radical e |cos x|\ng(x) = √[(x−2)(x−4)] / (1 – |cos x|)\nDomínio\nZeros")
    # 8
    add_question(doc, "8) Função por ramos\nf(x) = ln(1 − x) se x > 0\n       e^x − 1  se x ≤ 0\nEstudar continuidade em x = 0.")
    # 9
    add_question(doc, "9) Função quadrática\nf(x) = x² − 6x + 5\na) Zeros\nb) Vértice + eixo de simetria\nc) Concavidade")
    # 10
    add_question(doc, "10) Limites — forma 1\nL1 = lim (n→∞) [ √(n² + 8n + 1) − n ]\nL2 = lim (n→∞) [ √(n² + 12n + 16) − √(n² + 4n + 4) ]\nL3 = lim (n→∞) [ √(3n + 4) − √(3n + 1) ]")

    # Q11-Q23 (continuação das perguntas)
    add_question(doc, "11) Quadrado Perfeito\nlim (x→2) (x² − 4x + 4) / (x² − 2x)")
    add_question(doc, "12) Diferença de Cubos\nlim (x→1) (x³ − 1) / (x⁴ − 1)")
    add_question(doc, "13) Grau 2 / Grau 4\nlim (x→∞) (4x² − x + 1) / (x⁴ + x³ − 2)")
    add_question(doc, "14) Derivadas básicas\na) 5x⁴\nb) x² − 3x\nc) 7/x³\nd) √(3x + 2)")
    add_question(doc, "15) Derivada pela definição\nf(x) = x²\nCalcular f’(x) pela definição.")
    add_question(doc, "16) Derivadas (cadeia, produto, quociente)\na) e^{x cos x}\nb) (sin(2x)) / x²")
    add_question(doc, "17) Produto\nh(x) = x³ cos x")
    add_question(doc, "18) Quociente\ng(x) = x² / (x + 1)")
    add_question(doc, "19) Regra da Cadeia\nk(x) = √(5x + 1)")
    add_question(doc, "20) Derivada + tangente\nf(x) = x² e^x (pede derivada e equação da tangente no ponto x=1)")
    add_question(doc, "21) Limite racional\nlim (x→−1) (x² − 1) / (x³ + 1)")
    add_question(doc, "22) Limite trigonométrico\nlim (x→0) (tan x) / x")
    add_question(doc, "23) Limite trigonométrico (1 − cos)\nlim (x→0) (1 − cos(5x)) / x")

    # Separa e passa soluções para a mesma página (ou seguinte)
    doc.add_page_break()
    add_header_paragraph(doc, "SOLUÇÕES — Versão 43 (Respostas rápidas)")

    # Soluções compactas Q1-Q23
    sol_lines = [
        "Q1 — C(n+3,3) – C(n+2,3) = C(n+2,2) -> (n+2)(n+1)/2 = 12(n+1) -> n = 22",
        "Q2 — a) P(2 damas) = C(4,2)/C(52,2) = 6/1326 = 1/221; b) P(naipes diferentes) = (4*13*13)/1326 = 13/17; c) P(1 figura OU carta preta) = (12+26-6)/52 = 8/13",
        "Q3 — Exemplo N=180: ni (dados) = 28, 40, 46, 66 -> f_i = n_i/180; N_i acumulada; F_i = N_i/180; Média = (1*28+2*40+3*46+4*66)/180 ≈ 2.95; Moda = classe 3; Mediana = classe 3",
        "Q4 — a) P(soma=9) = 4/36 = 1/9; b) P(pelo menos um 4) = 1-(5/6)^2 = 11/36; c) P(preta OU figura) = (26+12-6)/52 = 8/13",
        "Q5 — a_{n+1} = (2n+9)/(n+6); sequência tende para 2 (limite); classificação: eventualmente crescente para n grandes",
        "Q6 — Domínio: e^x - 2 > 0 -> x > ln(2); Zero: e^x = 3 -> x = ln 3",
        "Q7 — Domínio: x ≤ 2 ou x ≥ 4, e |cos x| ≠ 1; Zeros: x=2, x=4",
        "Q8 — Lim x→0^+ ln(1-x)=0; Lim x→0^- e^x-1=0 -> contínua em 0",
        "Q9 — f(x)=x^2-6x+5 -> zeros x=1 e x=5; vértice (3,-4); concavidade para cima",
        "Q10 — L1=4; L2=4; L3=0",
        "Q11 — Lim=0",
        "Q12 — Resultado = 3/4",
        "Q13 — Lim=0",
        "Q14 — (5x^4)'=20x^3; (x^2-3x)'=2x-3; (7/x^3)'=-21/x^4; (√(3x+2))'=3/(2√(3x+2))",
        "Q15 — f'(x)=2x (pela definição)",
        "Q16 — a) e^{x cos x}(cos x - x sin x); b) 2(x cos(2x)-sin(2x))/x^3",
        "Q17 — 3x^2 cos x - x^3 sin x",
        "Q18 — (x^2/(x+1))' = (x^2 + 2x)/(x+1)^2",
        "Q19 — 5/(2√(5x+1))",
        "Q20 — f(1)=e; f'(x)=e^x(x^2+2x) -> f'(1)=3e -> tangent: y = 3e x - 2e",
        "Q21 — -2/3",
        "Q22 — 1",
        "Q23 — 0"
    ]

    for line in sol_lines:
        add_solution_line(doc, line)

    # salvar ficheiro no Desktop do utilizador ruiferreira
    try:
        doc.save(fullpath)
        print(f"✅ Ficheiro gerado em: {fullpath}")
    except Exception as e:
        print("❌ Erro ao gravar ficheiro:", e)
        print("Tenta gravar num caminho válido do teu sistema, por exemplo: /Users/teu_user/Desktop/arquivo.docx")

if __name__ == "__main__":
    main()