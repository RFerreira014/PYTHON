# gera_parteA_v43.py
# Gera "Teste Diário — M23 Engenharia Informática — Versão 43 — Parte A (Enunciado Q1–Q10)"
# Requisitos: python3, pip install python-docx

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def set_normal_font(doc, name='Helvetica', size=11):
    style = doc.styles['Normal']
    style.font.name = name
    # Ensure font works across platforms
    style._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    style.font.size = Pt(size)

def create_partA_docx(path_out="Teste_Diario_M23_V43_ParteA.docx"):
    doc = Document()
    set_normal_font(doc, name='Helvetica', size=11)

    # Set page margins (2 cm)
    sections = doc.sections
    for s in sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2)
        s.right_margin = Cm(2)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Teste Diário — M23 Engenharia Informática — Versão 43 — Parte A (Enunciado)")
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph("")

    # ENUNCIADO header
    h = doc.add_paragraph("ENUNCIADO")
    h.runs[0].bold = True

    # Q1
    doc.add_paragraph("1) Identidade Combinatória (diferença — sem simetria feita)")
    doc.add_paragraph("Resolve em n:")
    doc.add_paragraph("C(n+3, 3) – C(n+2, 3) = 12(n+1)")

    # Q2
    doc.add_paragraph("")
    doc.add_paragraph("2) Probabilidades — baralho de 52 cartas")
    doc.add_paragraph("Tiram-se 2 cartas sem reposição.")
    doc.add_paragraph("a) P(2 damas)")
    doc.add_paragraph("b) P(2 cartas de naipes diferentes)")
    doc.add_paragraph("c) P(1 figura OU carta preta)")

    # Q3 Estatística with horizontal compact table (N=180 for V43 - but per your request earlier this is Q1–Q10 set,
    # here we use N=180 style as earlier plan; adjust N if needed)
    doc.add_paragraph("")
    doc.add_paragraph("3) Estatística — Tabela (N = 180)")
    doc.add_paragraph("Preenche todos os campos e calcula média, mediana e moda.")
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'  # minimal grid; Pages will accept it fine

    headers = ["Classe", "nᵢ", "fᵢ", "Nᵢ", "Fᵢ"]
    hdr_cells = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr_cells[i].text = htext

    # Data salteada: some nᵢ and some fᵢ given, rest blank to fill by hand
    data = [
        ["1", "28", "", "", ""],
        ["2", "", "0,22", "", ""],
        ["3", "46", "", "", ""],
        ["4", "", "0,18", "", ""],
        ["Soma", "180", "", "", "1"]
    ]
    for r, row in enumerate(data, start=1):
        cells = table.rows[r].cells
        for c, val in enumerate(row):
            cells[c].text = val

    # center align table cells
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Observações: última classe ajusta para N=180.")
    doc.add_paragraph("")

    # Q4
    doc.add_paragraph("4) Probabilidades (dados e cartas)")
    doc.add_paragraph("a) 2 dados: P(soma = 9)")
    doc.add_paragraph("b) 2 dados: P(pelo menos um 4)")
    doc.add_paragraph("c) Baralho: P(preta OU figura)")

    # Q5
    doc.add_paragraph("")
    doc.add_paragraph("5) Sucessão racional")
    doc.add_paragraph("aₙ = (2n + 7) / (n + 5)")
    doc.add_paragraph("a) Calcula aₙ₊₁")
    doc.add_paragraph("b) Estuda o sinal de (aₙ₊₁ − aₙ)")
    doc.add_paragraph("c) Limite")
    doc.add_paragraph("d) Classificação")

    # Q6
    doc.add_paragraph("")
    doc.add_paragraph("6) Função logarítmica")
    doc.add_paragraph("f(x) = ln( 1 / (e^x − 2) )")
    doc.add_paragraph("a) Domínio")
    doc.add_paragraph("b) Zeros")

    # Q7
    doc.add_paragraph("")
    doc.add_paragraph("7) Radical e |cos x|")
    doc.add_paragraph("g(x) = sqrt[(x − 2)(x − 4)] / (1 − |cos x|)")
    doc.add_paragraph("Domínio")
    doc.add_paragraph("Zeros")

    # Q8
    doc.add_paragraph("")
    doc.add_paragraph("8) Função por ramos")
    doc.add_paragraph("f(x) =")
    doc.add_paragraph("  ln(1 − x)    se x > 0")
    doc.add_paragraph("  e^x − 1      se x ≤ 0")
    doc.add_paragraph("Estudar continuidade em x = 0.")

    # Q9
    doc.add_paragraph("")
    doc.add_paragraph("9) Função quadrática")
    doc.add_paragraph("f(x) = x^2 − 6x + 5")
    doc.add_paragraph("a) Zeros")
    doc.add_paragraph("b) Vértice + eixo de simetria")
    doc.add_paragraph("c) Concavidade")

    # Q10
    doc.add_paragraph("")
    doc.add_paragraph("10) Limites — forma 1")
    doc.add_paragraph("L1 = lim_{n→∞} [ sqrt(n^2 + 8n + 1) − n ]")
    doc.add_paragraph("L2 = lim_{n→∞} [ sqrt(n^2 + 12n + 16) − sqrt(n^2 + 4n + 4) ]")
    doc.add_paragraph("L3 = lim_{n→∞} [ sqrt(3n + 4) − sqrt(3n + 1) ]")

    # Save
    doc.save(path_outteste)
    print(f"Ficheiro gerado: {os.path.abspath(path_out)}")

if __name__ == "__main__":
    create_partA_docx()
  doc.save("/Users/rui/Desktop/Teste_Diario_M23_V43_ParteB_Solucoes.docx")
print("✅ Ficheiro gravado no Ambiente de Trabalho!")
doc.save("/Users/rui/Desktop/Teste_Diario_M23_V43_ParteB_Solucoes.docx")
print("✅ Ficheiro gravado no Ambiente de Trabalho!")